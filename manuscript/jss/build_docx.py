"""Build an editable Word reading copy from the authoritative JSS LaTeX file."""

from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "faissR_jss.tex"
OUTPUT = HERE / "faissR_jss.docx"
REFERENCE = HERE / "faissR_jss.docx"


def extract_braced_command(text: str, command: str) -> str:
    marker = f"\\{command}{{"
    start = text.index(marker) + len(marker)
    depth = 1
    pos = start
    while depth and pos < len(text):
        if text[pos] == "{" and (pos == 0 or text[pos - 1] != "\\"):
            depth += 1
        elif text[pos] == "}" and (pos == 0 or text[pos - 1] != "\\"):
            depth -= 1
        pos += 1
    if depth:
        raise ValueError(f"Unbalanced command: {command}")
    return text[start : pos - 1]


def normalize_body(body: str) -> str:
    cross_references = {
        "tab:api": ("Table", "1"),
        "tab:methods": ("Table", "2"),
        "tab:tuninggrid": ("Table", "3"),
        "tab:calibrationaudit": ("Table", "4"),
        "tab:datasets": ("Table", "4"),
        "tab:interim-heldout": ("Table", "5"),
        "tab:evidenceaudit": ("Table", "6"),
        "tab:ablations": ("Table", "7"),
        "sec:evaluation": ("Section", "7"),
    }
    for label, (kind, number) in cross_references.items():
        body = body.replace(
            f"{kind}~\\ref{{{label}}}",
            f"{kind} {number}",
        )

    def replace_path(match: re.Match[str]) -> str:
        path = match.group(1).replace("_", r"\_")
        return rf"\texttt{{{path}}}"

    body = re.sub(r"\\path\{([^}]+)\}", replace_path, body)

    replacements = {
        "pkg": "texttt",
        "code": "texttt",
        "class": "texttt",
        "fct": "texttt",
        "method": "texttt",
        "proglang": "textsf",
        "email": "texttt",
    }
    for source, target in replacements.items():
        body = body.replace(f"\\{source}{{", f"\\{target}{{")

    body = body.replace("\\begin{CodeChunk}\n", "")
    body = body.replace("\\end{CodeChunk}\n", "")
    body = body.replace("\\begin{CodeInput}", "\\begin{verbatim}")
    body = body.replace("\\end{CodeInput}", "\\end{verbatim}")
    body = body.replace("\\begin{CodeOutput}", "\\begin{verbatim}")
    body = body.replace("\\end{CodeOutput}", "\\end{verbatim}")
    body = re.sub(r"\\label\{[^}]+\}", "", body)
    body = re.sub(r"\\bibliography\{[^}]+\}", "", body)
    return body


def draw_flow(path: Path, lines: list[str], bold_line: int | None = None) -> None:
    width = 1600
    padding = 70
    row_height = 110
    height = padding * 2 + row_height * len(lines)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    regular_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    regular = ImageFont.truetype(regular_path, 34)
    bold = ImageFont.truetype(bold_path, 34)
    draw.rounded_rectangle(
        (15, 15, width - 15, height - 15), radius=8, outline="black", width=3
    )
    for index, line in enumerate(lines):
        font = bold if index == bold_line else regular
        box = draw.textbbox((0, 0), line, font=font)
        x = (width - (box[2] - box[0])) / 2
        y = padding + index * row_height
        draw.text((x, y), line, fill="black", font=font)
        if index < len(lines) - 1:
            arrow_x = width / 2
            arrow_top = y + 52
            arrow_bottom = y + row_height - 10
            draw.line((arrow_x, arrow_top, arrow_x, arrow_bottom), fill="black", width=3)
            draw.polygon(
                [
                    (arrow_x - 10, arrow_bottom - 12),
                    (arrow_x + 10, arrow_bottom - 12),
                    (arrow_x, arrow_bottom),
                ],
                fill="black",
            )
    image.save(path, dpi=(220, 220))


def replace_flow_figures(body: str, architecture: Path, validation: Path) -> str:
    architecture_figure = rf"""\begin{{figure}}[t!]
\centering
\includegraphics[width=0.92\linewidth]{{{architecture.as_posix()}}}
\caption{{Execution architecture. Backend resolution, metric transformation,
provider execution, and result residency are separate recorded decisions.}}
\end{{figure}}"""
    validation_figure = rf"""\begin{{figure}}[t!]
\centering
\includegraphics[width=0.92\linewidth]{{{validation.as_posix()}}}
\caption{{Calibration and validation workflow. Independence is at the sampled
query level within the same datasets, not at the dataset-collection level.}}
\end{{figure}}"""
    body = re.sub(
        r"\\begin\{figure\}\[t!\](?:(?!\\end\{figure\}).)*?Calibration query subset(?:(?!\\end\{figure\}).)*?\\end\{figure\}",
        lambda _: validation_figure,
        body,
        count=1,
        flags=re.DOTALL,
    )
    body = re.sub(
        r"\\begin\{figure\}\[t!\](?:(?!\\end\{figure\}).)*?Execution architecture\.(?:(?!\\end\{figure\}).)*?\\end\{figure\}",
        lambda _: architecture_figure,
        body,
        count=1,
        flags=re.DOTALL,
    )
    return body


def build_intermediate(source: str, architecture: Path, validation: Path) -> str:
    abstract = extract_braced_command(source, "Abstract")
    keywords = extract_braced_command(source, "Keywords")
    body = source.split("\\begin{document}", 1)[1]
    body = body.rsplit("\\end{document}", 1)[0]
    body = replace_flow_figures(body, architecture, validation)
    body = normalize_body(body)

    title = "faissR: Recall-Aware CPU and GPU Nearest-Neighbor Search in R"
    authors = (
        "Moussa Kassim (1,2; co-first); Martin Ocharo (1,2; co-first); "
        "Dalia Ahmed (1); Dupe Ojo (1); Alessia Vignoli (3,4); "
        "Leonardo Tenori (3,4; co-corresponding); "
        "Dinesh Gupta (20); Silvano Piazza (21,22); "
        "Stefano Cacciatore (1,2; co-corresponding)"
    )
    front = f"""\\documentclass{{article}}
\\usepackage{{booktabs}}
\\usepackage{{array}}
\\usepackage{{tabularx}}
\\usepackage{{amsmath}}
\\usepackage{{amssymb}}
\\title{{{title}}}
\\author{{{authors}}}
\\begin{{document}}
\\maketitle
\\begin{{abstract}}
{normalize_body(abstract)}
\\end{{abstract}}
\\paragraph{{Keywords}} {normalize_body(keywords)}
\\section*{{Correspondence and affiliations}}
1. Bioinformatics Unit, International Centre for Genetic Engineering and
Biotechnology (ICGEB), Cape Town 7925, South Africa.

2. Department of Integrative Biomedical Sciences, Institute of Infectious
Disease \\& Molecular Medicine (IDM), University of Cape Town, Cape Town 7925,
South Africa.

3. Department of Chemistry ``Ugo Schiff'', University of Florence, Sesto
Fiorentino, Italy.

4. Magnetic Resonance Center (CERM), University of Florence, Sesto Fiorentino,
Italy.

Moussa Kassim and Martin Ocharo contributed equally and share first authorship.

Co-corresponding authors:

Leonardo Tenori, \\texttt{{TENORI@CERM.UNIFI.IT}}

Stefano Cacciatore, \\texttt{{STEFANO.CACCIATORE@ICGEB.ORG}}

Author e-mails:

Moussa Kassim, \\texttt{{MOUSSA.KASSIM@ICGEB.ORG}}

Martin Ocharo, \\texttt{{MARTIN.OCHARO@ICGEB.ORG}}

Dalia Ahmed, \\texttt{{DALIA.AHMED@ICGEB.ORG}}

Dupe Ojo, \\texttt{{DUPE.OJO@ICGEB.ORG}}

Alessia Vignoli, \\texttt{{VIGNOLI@CERM.UNIFI.IT}}

Leonardo Tenori, \\texttt{{TENORI@CERM.UNIFI.IT}}

Stefano Cacciatore, \\texttt{{STEFANO.CACCIATORE@ICGEB.ORG}}

Stefano Cacciatore ORCID:
\\url{{https://orcid.org/0000-0001-7052-7156}}.
Project: \\url{{https://github.com/tkcaccia/faissR}}.
"""
    return front + body + "\n\\end{document}\n"


def polish_docx(path: Path) -> None:
    document = Document(path)
    in_correspondence = False
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        if stripped == "Correspondence and affiliations":
            in_correspondence = True
        elif stripped == "Introduction":
            in_correspondence = False
        if in_correspondence:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(3)
        if stripped.startswith(("R> ", "+ ", "[1]")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
    for table in document.tables:
        if not table.rows:
            continue
        row_properties = table.rows[0]._tr.get_or_add_trPr()
        if row_properties.find(qn("w:tblHeader")) is None:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_widths = {
            ("Evidence stream", "Archive status", "Consequence for this draft"):
                [2150, 2750, 4460],
            ("Ablation", "Backend", "Primary contrast", "Reported estimand"):
                [1700, 1200, 2980, 3480],
            ("Backend", "Method", "Metric", "Data", "Cells",
             "Quality status", "Median (s)"):
                [1100, 1550, 1550, 1100, 700, 2200, 1000],
        }
        widths = table_widths.get(tuple(headers))
        if widths is not None:
            table.autofit = False
            table_properties = table._tbl.tblPr
            layout = table_properties.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                table_properties.append(layout)
            layout.set(qn("w:type"), "fixed")
            table_width = table_properties.find(qn("w:tblW"))
            table_width.set(qn("w:type"), "dxa")
            table_width.set(qn("w:w"), str(sum(widths)))
            table_grid = table._tbl.tblGrid
            for grid_column in list(table_grid):
                table_grid.remove(grid_column)
            for width in widths:
                grid_column = OxmlElement("w:gridCol")
                grid_column.set(qn("w:w"), str(width))
                table_grid.append(grid_column)
            for row in table.rows:
                for cell, width in zip(row.cells, widths):
                    cell_properties = cell._tc.get_or_add_tcPr()
                    cell_width = cell_properties.find(qn("w:tcW"))
                    if cell_width is None:
                        cell_width = OxmlElement("w:tcW")
                        cell_properties.append(cell_width)
                    cell_width.set(qn("w:type"), "dxa")
                    cell_width.set(qn("w:w"), str(width))
    document.save(path)


def main() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    with tempfile.TemporaryDirectory(prefix="faissR-jss-docx-") as tmp:
        tmp_path = Path(tmp)
        architecture = tmp_path / "architecture.png"
        validation = tmp_path / "validation.png"
        draw_flow(
            architecture,
            [
                "R double matrix or optional float matrix",
                "float32 matrix view and metric transformation/cache",
                "compiled backend, method, and parameter selection",
                "FAISS CPU | FAISS GPU | cuVS/CUDA | native compiled route",
                "host result or explicit GPU-resident exact-family object",
            ],
        )
        draw_flow(
            validation,
            [
                "Calibration query subset (seed 4)",
                "method-specific grids and failure filtering",
                "compiled static policies; validation cannot change selection",
                "independent query subsets (seeds 20260706 and 20260807)",
                "three timing repeats per seed",
                "completion, recall, auto-versus-oracle, and ablation summaries",
            ],
            bold_line=2,
        )
        intermediate = Path(tmp) / "faissR_jss_word.tex"
        intermediate.write_text(
            build_intermediate(source, architecture, validation), encoding="utf-8"
        )
        command = [
            "pandoc",
            str(intermediate),
            "--from=latex",
            "--to=docx",
            "--standalone",
            "--citeproc",
            f"--bibliography={HERE / 'faissR_jss.bib'}",
            f"--resource-path={HERE}",
            f"--output={OUTPUT}",
        ]
        if REFERENCE.exists():
            reference_copy = Path(tmp) / "reference.docx"
            reference_copy.write_bytes(REFERENCE.read_bytes())
            command.append(f"--reference-doc={reference_copy}")
        subprocess.run(command, check=True, cwd=HERE)
        polish_docx(OUTPUT)


if __name__ == "__main__":
    main()
