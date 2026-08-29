"""Build an editable Word copy of the JSS supplementary material."""

from pathlib import Path
import re
import subprocess
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "faissR_jss_supplement.tex"
OUTPUT = HERE / "faissR_jss_supplement.docx"
REFERENCE = HERE / "faissR_jss.docx"


def word_source(source: str) -> str:
    """Normalize LaTeX constructs that Pandoc does not map cleanly to Word."""
    source = re.sub(
        r"\\path\{([^{}]+)\}",
        lambda match: r"\texttt{" + match.group(1).replace("_", r"\_") + "}",
        source,
    )
    source = source.replace(r"\textsuperscript{\(\dagger\)}", " ")
    source = source.replace(r"\dagger", "")
    source = source.replace(r"\ast", "")
    source = re.sub(r"\$\^\{([^}]+)\}\$", r" (\1)", source)
    source = source.replace(r"\newcolumntype{Y}{>{\raggedright\arraybackslash}X}", "")
    source = source.replace(
        r"\newcolumntype{P}[1]{>{\raggedright\arraybackslash}p{#1}}",
        "",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{p{0.27\linewidth}Y}",
        r"\begin{tabular}{p{0.27\linewidth}p{0.63\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\textwidth}{@{}>{\raggedright\arraybackslash}p{0.42\textwidth}Y@{}}",
        r"\begin{tabular}{p{0.42\textwidth}p{0.48\textwidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{P{0.30\linewidth}Y}",
        r"\begin{tabular}{p{0.30\linewidth}p{0.60\linewidth}}",
    )
    source = source.replace(
        "\\begin{tabularx}{\\linewidth}{P{0.19\\linewidth}"
        "P{0.24\\linewidth}Y\nP{0.13\\linewidth}}",
        "\\begin{tabular}{p{0.19\\linewidth}p{0.24\\linewidth}"
        "p{0.31\\linewidth}p{0.13\\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{lrrY}",
        r"\begin{tabular}{lrrp{0.55\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{p{0.32\linewidth}p{0.18\linewidth}Y}",
        r"\begin{tabular}{p{0.32\linewidth}p{0.18\linewidth}p{0.40\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{p{0.25\linewidth}Y}",
        r"\begin{tabular}{p{0.25\linewidth}p{0.65\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{P{0.25\linewidth}Y}",
        r"\begin{tabular}{p{0.25\linewidth}p{0.65\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\linewidth}{P{0.32\linewidth}P{0.18\linewidth}Y}",
        r"\begin{tabular}{p{0.32\linewidth}p{0.18\linewidth}p{0.40\linewidth}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\textwidth}{@{}lY@{}}",
        r"\begin{tabular}{@{}lp{0.72\textwidth}@{}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\textwidth}{@{}Yrr@{}}",
        r"\begin{tabular}{@{}p{0.62\textwidth}rr@{}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\textwidth}{@{}Xrr@{}}",
        r"\begin{tabular}{@{}p{0.62\textwidth}rr@{}}",
    )
    source = source.replace(
        r"\begin{tabularx}{\textwidth}{@{}Xrrrr@{}}",
        r"\begin{tabular}{@{}p{0.42\textwidth}rrrr@{}}",
    )
    source = source.replace(
        r"\begin{longtable}{P{0.22\linewidth}P{0.31\linewidth}P{0.37\linewidth}}",
        r"\begin{longtable}{p{0.22\linewidth}p{0.31\linewidth}p{0.37\linewidth}}",
    )
    source = source.replace(
        r"\begin{longtable}{P{0.18\linewidth}P{0.75\linewidth}}",
        r"\begin{longtable}{p{0.18\linewidth}p{0.75\linewidth}}",
    )
    source = source.replace(
        r"\begin{longtable}{P{0.20\linewidth}P{0.72\linewidth}}",
        r"\begin{longtable}{p{0.20\linewidth}p{0.72\linewidth}}",
    )
    source = source.replace(r"\end{tabularx}", r"\end{tabular}")
    source = re.sub(
        r"\\endfirsthead.*?\\endhead",
        "",
        source,
        flags=re.DOTALL,
    )
    return source


def polish(path: Path) -> None:
    document = Document(path)
    for section in document.sections:
        section.header_distance = Pt(35.4)
        section.footer_distance = Pt(35.4)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name in {"First Paragraph", "Body Text"} and (
            text.startswith("(1) Bioinformatics Unit")
            or text.startswith("Moussa Kassim and Martin Ocharo")
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(4)
        if text == "Purpose and evidence boundary":
            paragraph.paragraph_format.page_break_before = True
        if text.startswith("Approximate method"):
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(42)
        if text.startswith("The datasets are COIL20"):
            page_break = paragraph.insert_paragraph_before()
            page_break.add_run().add_break(WD_BREAK.PAGE)
        if text.startswith(("CPU IVF", "CUDA IVF")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for table in document.tables:
        if not table.rows:
            continue
        while table.rows and all(
            not cell.text.strip() for cell in table.rows[0].cells
        ):
            table._tbl.remove(table.rows[0]._tr)
        if not table.rows:
            continue
        headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
        compact = headers == ("Backend/method", "Metric", "Recall at 15")
        table.autofit = False
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        if header_properties.find(qn("w:tblHeader")) is None:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            header_properties.append(repeat)
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:cantSplit")) is None:
                cant_split = OxmlElement("w:cantSplit")
                cant_split.set(qn("w:val"), "true")
                row_properties.append(cant_split)
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                margins = tc_pr.find(qn("w:tcMar"))
                if margins is None:
                    margins = OxmlElement("w:tcMar")
                    tc_pr.append(margins)
                vertical_margin = 25 if compact else 80
                for side, value in (
                    ("top", vertical_margin),
                    ("bottom", vertical_margin),
                    ("start", 120),
                    ("end", 120),
                ):
                    node = margins.find(qn(f"w:{side}"))
                    if node is None:
                        node = OxmlElement(f"w:{side}")
                        margins.append(node)
                    node.set(qn("w:w"), str(value))
                    node.set(qn("w:type"), "dxa")
                if compact:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.size = Pt(8.5)
        table_widths = {
            ("Component", "Configuration"): [2400, 6960],
            ("Method", "Tuned quantities"): [2500, 6860],
            ("Evidence", "Passing or completed", "Total"):
                [5600, 2160, 1600],
            ("Function", "Role"): [2500, 6860],
            ("Method", "CPU route", "CUDA route"): [1600, 3380, 4380],
            ("Dataset", "Rows", "Columns"): [4000, 2680, 2680],
            (
                "Dataset",
                "Source/release",
                "Representation searched",
                "Terms",
            ): [1800, 2200, 4060, 1300],
            (
                "Backend",
                "Contract pass",
                "Unsupported",
                "Target 0.99 pass",
                "Edge-case pass",
            ): [1450, 1900, 1550, 2100, 2360],
            ("Backend/method", "Metric", "Recall at 15"): [3500, 3000, 2860],
            ("Method", "Candidate settings"): [1900, 7460],
            ("Package", "Public interface and timing interpretation"):
                [1900, 7460],
            ("Requested names", "Publication treatment"): [2500, 6860],
            ("Backend", "Rows", "Datasets", "Covered contrasts"):
                [1200, 900, 1200, 6060],
            ("Evidence stream", "Status", "Required action"): [3000, 1800, 4560],
        }
        widths = table_widths.get(headers)
        if widths is not None:
            table_properties = table._tbl.tblPr
            layout = table_properties.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                table_properties.append(layout)
            layout.set(qn("w:type"), "fixed")
            table_width = table_properties.find(qn("w:tblW"))
            table_width.set(qn("w:type"), "dxa")
            table_width.set(qn("w:w"), str(sum(widths)))
            table_grid = table._tbl.tblGrid
            for grid_column in list(table_grid):
                table_grid.remove(grid_column)
            for width in widths:
                grid_column = OxmlElement("w:gridCol")
                grid_column.set(qn("w:w"), str(width))
                table_grid.append(grid_column)
            for row in table.rows:
                for cell, width in zip(row.cells, widths):
                    cell_properties = cell._tc.get_or_add_tcPr()
                    cell_width = cell_properties.find(qn("w:tcW"))
                    if cell_width is None:
                        cell_width = OxmlElement("w:tcW")
                        cell_properties.append(cell_width)
                    cell_width.set(qn("w:type"), "dxa")
                    cell_width.set(qn("w:w"), str(width))
    document.save(path)


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="faissR-jss-supplement-") as tmp:
        source = Path(tmp) / "supplement.tex"
        source.write_text(
            word_source(SOURCE.read_text(encoding="utf-8")),
            encoding="utf-8",
        )
        intermediate = Path(tmp) / "supplement.docx"
        command = [
            "pandoc",
            str(source),
            "--from=latex",
            "--to=docx",
            "--standalone",
            f"--output={intermediate}",
        ]
        if REFERENCE.exists():
            command.append(f"--reference-doc={REFERENCE}")
        subprocess.run(command, check=True, cwd=HERE)
        OUTPUT.write_bytes(intermediate.read_bytes())
    polish(OUTPUT)


if __name__ == "__main__":
    main()
